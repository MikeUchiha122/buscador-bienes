from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('GUIA COMPLETA DEL BUSCADOR DE BIENES SICOBIM', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Explicacion para ninos de 5 anos (y adultos que empiezan)')
doc.add_paragraph('')

doc.add_heading('1. QUIEN ES ESTA APLICACION?', level=1)
doc.add_paragraph('Imagina que tienes una biblioteca con 1000 libros y necesitas encontrar uno rapido. Esta aplicacion es como un bibliotecario inteligente que te ayuda a encontrar cualquier bien en una lista enorme.')
doc.add_paragraph('Se llama Buscador de Bienes SICOBIM y esta hecho para el ISSSTE.')

doc.add_heading('2. PARA QUE SIRVE?', level=1)
funciones = [
    'Buscar bienes por su numero',
    'Buscar por descripcion',
    'Buscar por ubicacion',
    'Ver el estado del bien (bueno, malo, regular)',
    'Editar informacion de los bienes',
    'Dar de baja un bien',
    'Cruzar informacion con otros archivos Excel'
]
for f in funciones:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('3. QUE HERRAMIENTAS USA? (FRAMEWORKS)', level=1)

doc.add_heading('Tkinter', level=2)
doc.add_paragraph('Es como los bloques LEGO para hacer ventanas y botones.')

doc.add_heading('Pandas', level=2)
doc.add_paragraph('Es como una hoja de calculo muy inteligente en el codigo.')

doc.add_heading('OpenPyXL', level=2)
doc.add_paragraph('Es el traductor entre Python y Excel.')

doc.add_heading('Python', level=2)
doc.add_paragraph('El lenguaje de programacion principal.')

doc.add_heading('4. ESTRUCTURA', level=1)
doc.add_paragraph('bienes-python/')
doc.add_paragraph('  buscador_bienes.py        (Codigo principal)')
doc.add_paragraph('  BASE_BIENES_UNIDAD.xlsx  (Base de datos)')
doc.add_paragraph('  archivo_cruzar/          (Archivos a cruzar)')
doc.add_paragraph('  bienes.log               (Registro de acciones)')

doc.add_heading('5. CODIGO EXPLICADO', level=1)

doc.add_heading('5.1 Importar herramientas', level=2)
doc.add_paragraph('Antes de construir algo, necesitamos las herramientas:')
doc.add_paragraph('- tkinter: Para hacer ventanas y botones')
doc.add_paragraph('- pandas: Para manejar hojas de calculo')
doc.add_paragraph('- os: Para manejar archivos')
doc.add_paragraph('- logging: Cuaderno de notas')

doc.add_heading('5.2 Configuracion', level=2)
doc.add_paragraph('El programa necesita saber donde esta instalado, donde escribir el diario y donde esta el Excel.')

doc.add_heading('5.3 Clase principal', level=2)
doc.add_paragraph('class BuscadorBienes: Es como un molde. Todos los objetos creados tienen las mismas caracteristicas.')
doc.add_paragraph('- self.root: La ventana principal')
doc.add_paragraph('- .geometry(1200x700): Tamano de la ventana')
doc.add_paragraph('- .configure(bg=#1e1e2e): Color de fondo')

doc.add_heading('5.4 Cargar datos', level=2)
doc.add_paragraph('1. Lee el archivo Excel')
doc.add_paragraph('2. Limpia los nombres de las columnas')
doc.add_paragraph('3. Convierte el numero de inventario a numero')
doc.add_paragraph('4. Si no existe la columna STATUS, la crea')
doc.add_paragraph('5. Guarda los cambios')

doc.add_heading('5.5 Interfaz', level=2)
doc.add_paragraph('Crea los elementos que ves en pantalla:')
doc.add_paragraph('- Frame: Contenedor')
doc.add_paragraph('- Label: Texto que se muestra')
doc.add_paragraph('- Entry: Caja para escribir')
doc.add_paragraph('- Button: Boton')
doc.add_paragraph('- Treeview: Tabla de resultados')
doc.add_paragraph('- Combobox: Menu desplegable')

doc.add_heading('5.6 Buscar', level=2)
doc.add_paragraph('1. Obtiene lo que el usuario escribio')
doc.add_paragraph('2. Aplica el filtro (VIGENTES/BAJAS/TODOS)')
doc.add_paragraph('3. Busca el texto en los registros')
doc.add_paragraph('4. Muestra los resultados')

doc.add_heading('5.7 Ver detalle', level=2)
doc.add_paragraph('1. Abre una nueva ventana al hacer doble clic')
doc.add_paragraph('2. Muestra todos los datos del bien')
doc.add_paragraph('3. Permite editar algunos campos')
doc.add_paragraph('4. Permite guardar o dar de baja')

doc.add_heading('5.8 Cruzar archivos', level=2)
doc.add_paragraph('CRUZAR POR INVENTARIO:')
doc.add_paragraph('- Compara solo por NUMERO DE INVENTARIO')
doc.add_paragraph('- Existentes: estan en ambos archivos')
doc.add_paragraph('- Faltantes: solo estan en el nuevo')
doc.add_paragraph('')
doc.add_paragraph('CRUZAR POR INVENTARIO + CABM:')
doc.add_paragraph('- Compara por NUMERO + CABM')
doc.add_paragraph('- Solo hace match si ambos coinciden')

doc.add_heading('6. COMO USAR', level=1)

doc.add_heading('Paso 1: Ejecutar', level=2)
doc.add_paragraph('1. Abre la carpeta bienes-python')
doc.add_paragraph('2. Ejecuta: python buscador_bienes.py')
doc.add_paragraph('3. Veras la ventana principal')

doc.add_heading('Paso 2: Buscar', level=2)
doc.add_paragraph('1. Escribe el texto a buscar')
doc.add_paragraph('2. Selecciona donde buscar')
doc.add_paragraph('3. Selecciona el filtro')
doc.add_paragraph('4. Presiona ENTER o Buscar')

doc.add_heading('Paso 3: Ver detalle', level=2)
doc.add_paragraph('1. Haz doble clic en un registro')
doc.add_paragraph('2. Edita los campos que necesites')
doc.add_paragraph('3. Guarda o da de baja')

doc.add_heading('Paso 4: Cruzar', level=2)
doc.add_paragraph('1. Coloca archivo en archivo_cruzar/ARCHIVO_NUEVO.xlsx')
doc.add_paragraph('2. Presionaboton x Inventario o x Inv+CABM')
doc.add_paragraph('3. Se genera informe Excel')

doc.add_heading('7. GLOSARIO', level=1)
terminos = [
    ('DataFrame', 'Tabla de datos como Excel en Python'),
    ('Columna', 'Linea vertical de datos'),
    ('Fila', 'Linea horizontal de datos'),
    ('Frame', 'Contenedor de interfaz'),
    ('Label', 'Texto que se muestra'),
    ('Entry', 'Caja para escribir texto'),
    ('Button', 'Boton que se presiona'),
    ('Treeview', 'Tabla con resultados'),
    ('Combobox', 'Menu desplegable'),
]
for term, meaning in terminos:
    doc.add_paragraph(f'{term}: {meaning}')

doc.add_heading('8. RESUMEN', level=1)
doc.add_paragraph('Python: El idioma de la computadora')
doc.add_paragraph('Pandas: Maneja datos como Excel')
doc.add_paragraph('Tkinter: Crea ventanas y botones')
doc.add_paragraph('OpenPyXL: Guarda archivos Excel')
doc.add_paragraph('')
doc.add_paragraph('La aplicacion: Lee Excel, muestra tabla, permite buscar/editar/baja y cruzar archivos.')

doc.add_heading('CREDITOS', level=1)
doc.add_paragraph('Desarrollado por: Ing. Miguel Angel Ramirez Galicia')
doc.add_paragraph('Para: ISSSTE - SICOBIM')
doc.add_paragraph('Fecha: 2025')

doc.save('C:/bienes-python/GUIA_BUSCADOR_BIENES.docx')
print('Word creado!')